from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
from uuid import UUID, uuid4
import io
import json
import logging

from pypdf import PdfReader
from PIL import Image
import pytesseract
from docx import Document

from src.services.llm import chat_completion, extract_document_data

logger = logging.getLogger(__name__)

router = APIRouter()


class CoberturaExtract(BaseModel):
    nome: str
    valorLimite: Optional[float] = None
    franquia: Optional[float] = None
    incluido: bool = True


class OrcamentoExtract(BaseModel):
    seguradoraNome: Optional[str] = None
    valorPremio: Optional[float] = None
    dataVigenciaInicio: Optional[str] = None
    dataVigenciaFim: Optional[str] = None
    formaPagamento: Optional[str] = None
    descontos: Optional[float] = None
    coberturas: List[CoberturaExtract] = []


class DocumentExtractionResponse(BaseModel):
    documento_id: Optional[UUID] = None
    tipo: str
    dados_extraidos: Dict[str, Any]
    chunks_created: int = 0
    status: str
    success: bool = True
    message: Optional[str] = None
    texto_extraido: Optional[str] = None


class DocumentAnalysisRequest(BaseModel):
    documento_id: UUID
    tipo_analise: str = "completa"


EXTRACTION_PROMPT_ORCAMENTO = """Voce e um especialista em seguros de condominio analisando uma cotacao/orcamento de seguradora.

Sua tarefa: extrair TODOS os dados estruturados do documento abaixo em formato JSON.

REGRAS CRITICAS:

1. **valorPremio**: SEMPRE o PREMIO TOTAL (com IOF/taxas), NUNCA o premio liquido.
   - Procure: "Total a Pagar", "Premio Total", "Valor Total", "Total Geral"
   - Se so encontrar "Premio Liquido", some com IOF: liquido + iof = total
   - Exemplo: Liquido R$ 4.683,17 + IOF R$ 345,63 = valorPremio: 5028.80

2. **formaPagamento**: A QUANTIDADE MAXIMA DE PARCELAS SEM JUROS no boleto.
   - Procure tabelas de "Opcoes de Pagamento", "Parcelamento", "Plano de Pagamento"
   - Identifique APENAS as opcoes SEM JUROS (juros = 0% ou "sem juros")
   - Pegue o NUMERO MAXIMO de parcelas sem juros disponivel
   - Formato: "Ate XXx sem juros no boleto" (ex: "Ate 06x sem juros no boleto")
   - HDI normalmente: 6x sem juros / Allianz: 6x / AXA: 6x / Tokio: 6x / Chubb: 10x
   - Se for so a vista, use "A vista"

3. **coberturas**: Lista DETALHADA de TODAS as coberturas contratadas, com:
   - nome: Nome exato como esta no PDF (ex: "Incendio, Raio e Explosao")
   - valorLimite: Limite Maximo de Indenizacao (LMI) - numerico
   - franquia: Valor da franquia em REAIS - numerico (ou null se "Sem Franquia")
     * IMPORTANTE: procure na MESMA LINHA ou linhas proximas a cobertura
     * "Sem franquia" ou "Nao ha" = null
     * "20% dos prejuizos com minimo de R$ 5.000" = 5000
     * "Franquia: R$ 1.500" = 1500
   - incluido: true (so liste as contratadas)

4. **dadosImovel**: Dados do imovel/condominio segurado:
   - tipoCondominio: Residencial/Comercial/Misto/etc
   - numeroElevadores: numerico
   - anoConstrucao: ano (ex: 2010) ou idade ("De 06 a 10 anos")
   - numeroUnidades: numerico
   - numeroBlocos: numerico
   - numeroPavimentos: numerico ou faixa ("Acima de 15 andares")
   - areaConstruida: m2 (numerico)
   - numeroFuncionarios: numerico
   - placaSolar: true/false
   - bensAoArLivre: true/false
   - bonus: texto (ex: "Classe 5", "Sem sinistro 5 anos")
   - protecionais: lista (["Extintores", "Hidrantes", "Alarme"])

FORMATO DO JSON DE RESPOSTA:
{{
  "seguradoraNome": "...",
  "valorPremio": 5028.80,
  "dataVigenciaInicio": "2025-12-12",
  "dataVigenciaFim": "2026-12-12",
  "formaPagamento": "Ate 06x sem juros no boleto",
  "coberturas": [
    {{"nome": "Incendio, Raio e Explosao", "valorLimite": 22421000.00, "franquia": null, "incluido": true}},
    {{"nome": "Danos Eletricos", "valorLimite": 100000.00, "franquia": 1500.00, "incluido": true}}
  ],
  "dadosImovel": {{
    "tipoCondominio": "Residencial",
    "numeroElevadores": 1,
    "anoConstrucao": "De 06 a 10 anos",
    "numeroUnidades": 80,
    "numeroBlocos": 1,
    "numeroPavimentos": 16,
    "areaConstruida": 6000,
    "numeroFuncionarios": null,
    "placaSolar": false,
    "bensAoArLivre": false,
    "bonus": null,
    "protecionais": ["Extintores"]
  }}
}}

IMPORTANTE:
- Extraia APENAS o que esta EXPLICITAMENTE no texto
- Use null para campos nao encontrados
- Valores monetarios: numericos sem R$, sem pontos de milhar, ponto como decimal
- Seja PRECISO com franquias - leia com atencao a tabela de coberturas
- Responda APENAS com o JSON, SEM markdown, SEM ```, SEM explicacoes

TEXTO DO DOCUMENTO:
{texto}
"""

EXTRACTION_PROMPT_CONDOMINIO = """Voce e um assistente especializado em extrair dados de condominios de QUALQUER tipo de documento.

O documento pode ser de qualquer tipo: convencao, ata de assembleia, habite-se, regimento interno,
alvara, AVCB, laudo tecnico, planta, contrato, apolice de seguro, orcamento, etc.

Analise o texto abaixo e extraia as seguintes informacoes do CONDOMINIO em formato JSON:

- nome: Nome completo do condominio (ex: "Condominio Residencial Solar das Palmeiras")
- cnpj: CNPJ do condominio (formato XX.XXX.XXX/XXXX-XX)
- endereco: Logradouro/Rua (ex: "Rua das Flores")
- numero: Numero do endereco
- bairro: Nome do bairro
- cidade: Nome da cidade
- estado: Sigla do estado (ex: "SP", "RJ")
- cep: CEP (formato XXXXX-XXX)
- areaConstruida: Area construida em metros quadrados (numerico)
- numeroUnidades: Quantidade de unidades/apartamentos (numerico)
- numeroBlocos: Quantidade de blocos/torres (numerico)
- numeroAndares: Quantidade de andares/pavimentos (numerico)
- numeroElevadores: Quantidade de elevadores (numerico)
- anoConstrucao: Ano de construcao (numerico, ex: 1998)
- tipoConstrucao: Tipo do condominio (RESIDENCIAL, COMERCIAL, MISTO, RESIDENCIAL_HORIZONTAL, etc.)
- seguradoraAtual: Nome da seguradora atual se mencionada
- vencimentoApolice: Data de vencimento da apolice atual (formato YYYY-MM-DD)

IMPORTANTE:
- Extraia APENAS informacoes que estao explicitamente no texto
- Use null para campos nao encontrados
- Para areas, use apenas o numero sem "m2"
- Em CONVENCOES procure: nome do condominio, CNPJ, endereco, numero de unidades
- Em ATAS procure: nome do condominio, datas, decisoes
- Em HABITE-SE procure: endereco, area construida, numero de pavimentos
- Em AVCB/Alvara procure: endereco, area, numero de pavimentos
- Em apolices/orcamentos procure: TODAS as informacoes acima

Responda APENAS com o JSON, sem explicacoes.

TEXTO DO DOCUMENTO:
{texto}
"""


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extrai texto de um PDF"""
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao ler PDF: {str(e)}")


async def _extract_structured_data(texto: str, tipo: str) -> Dict[str, Any]:
    """Helper to extract structured data from text using LLM."""
    texto_truncado = texto[:50000]

    if tipo == "condominio":
        prompt = EXTRACTION_PROMPT_CONDOMINIO.format(texto=texto_truncado)
    else:
        prompt = EXTRACTION_PROMPT_ORCAMENTO.format(texto=texto_truncado)

    raw_response = await chat_completion(
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=2000,
    )

    try:
        json_str = raw_response.strip()
        if json_str.startswith("```"):
            json_str = json_str.split("\n", 1)[1]
        if json_str.endswith("```"):
            json_str = json_str.rsplit("```", 1)[0]
        json_str = json_str.strip()

        return json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error in extraction. Raw response: {raw_response[:500]}")
        return {"erro": f"Falha ao interpretar resposta da IA: {str(e)}", "raw_response": raw_response[:200]}


@router.post("/extract", response_model=DocumentExtractionResponse)
async def extract_document(
    file: UploadFile = File(...),
    tipo: str = "orcamento",
    condominio_id: Optional[UUID] = None,
):
    """
    Extrai dados de um documento (PDF)

    Tipos suportados:
    - orcamento: Orcamento de seguro
    - apolice: Apolice de seguro
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nome do arquivo nao informado")

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF sao suportados")

    try:
        # Read file content
        content = await file.read()

        # Extract text from PDF
        texto = extract_text_from_pdf(content)

        if not texto.strip():
            raise HTTPException(status_code=400, detail="Nao foi possivel extrair texto do PDF")

        # Extract structured data if orcamento/apolice
        dados_extraidos = {}
        if tipo in ("orcamento", "apolice", "condominio"):
            dados_extraidos = await _extract_structured_data(texto, tipo)

        return DocumentExtractionResponse(
            documento_id=None,
            tipo=tipo,
            dados_extraidos=dados_extraidos,
            chunks_created=0,
            status="success",
            success=True,
            message="Dados extraidos com sucesso",
            texto_extraido=texto,
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar documento: {str(e)}")


class ClassifyRequest(BaseModel):
    texto: str
    nome_arquivo: Optional[str] = None


class ClassifyResponse(BaseModel):
    tipo: str
    confianca: float
    justificativa: Optional[str] = None


CLASSIFY_PROMPT = """Voce e um classificador de documentos de seguro de condominio.
Analise o texto abaixo e classifique o documento em UMA das seguintes categorias:

- APOLICE: Apolice de seguro (documento formal da seguradora com detalhes da cobertura vigente)
- ORCAMENTO: Orcamento ou proposta de seguro (cotacao com valores, coberturas oferecidas)
- CONDICOES_GERAIS: Condicoes gerais de seguro (regras, clausulas, definicoes gerais da seguradora)
- LAUDO_VISTORIA: Laudo de vistoria tecnica do imovel
- SINISTRO: Documento relacionado a sinistro (abertura, acompanhamento, laudo de sinistro)
- OUTRO: Documento que nao se encaixa nas categorias acima

Responda APENAS com JSON no formato:
{{"tipo": "CATEGORIA", "confianca": 0.95, "justificativa": "breve explicacao"}}

NOME DO ARQUIVO: {nome_arquivo}

TEXTO DO DOCUMENTO (primeiros 2000 caracteres):
{texto}
"""


@router.post("/classify", response_model=ClassifyResponse)
async def classify_document(request: ClassifyRequest):
    """
    Classifica automaticamente o tipo de um documento de seguro.
    """
    try:
        texto_truncado = request.texto[:2000]
        prompt = CLASSIFY_PROMPT.format(
            texto=texto_truncado,
            nome_arquivo=request.nome_arquivo or "desconhecido",
        )

        response = await chat_completion(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=200,
        )

        try:
            json_str = response.strip()
            if json_str.startswith("```"):
                json_str = json_str.split("\n", 1)[1]
            if json_str.endswith("```"):
                json_str = json_str.rsplit("```", 1)[0]
            json_str = json_str.strip()

            dados = json.loads(json_str)
            return ClassifyResponse(
                tipo=dados.get("tipo", "OUTRO"),
                confianca=float(dados.get("confianca", 0.5)),
                justificativa=dados.get("justificativa"),
            )
        except json.JSONDecodeError:
            return ClassifyResponse(tipo="OUTRO", confianca=0.3, justificativa="Falha na classificacao automatica")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao classificar documento: {str(e)}")


class ExtractTextRequest(BaseModel):
    texto: str
    tipo: str = "orcamento"


class ExtractTextResponse(BaseModel):
    tipo: str
    dados_extraidos: Dict[str, Any]
    success: bool = True
    message: Optional[str] = None


@router.post("/extract-text", response_model=ExtractTextResponse)
async def extract_from_text(request: ExtractTextRequest):
    """
    Extrai dados estruturados a partir de texto ja extraido do PDF.
    Usado pelo pipeline automatico (backend envia texto, nao arquivo).
    Usa extract_document_data (Claude) para extracao completa incluindo condominio_data.
    """
    try:
        # Use the more complete extraction from llm.py that includes condominio_data
        dados = await extract_document_data(request.texto, request.tipo.lower())

        return ExtractTextResponse(
            tipo=request.tipo,
            dados_extraidos=dados,
            success=True,
            message="Dados extraidos com sucesso via Claude",
        )

    except Exception as e:
        logger.error(f"Erro na extracao via Claude, tentando fallback: {e}")
        # Fallback to simpler extraction
        try:
            dados = await _extract_structured_data(request.texto, request.tipo.lower())
            return ExtractTextResponse(
                tipo=request.tipo,
                dados_extraidos=dados,
                success=True,
                message="Dados extraidos com fallback",
            )
        except Exception as fallback_error:
            raise HTTPException(status_code=500, detail=f"Erro ao extrair dados: {str(fallback_error)}")


@router.post("/analyze")
async def analyze_document(request: DocumentAnalysisRequest):
    """
    Analisa um documento ja extraido
    """
    raise HTTPException(status_code=501, detail="Analise em desenvolvimento")


@router.get("/{documento_id}/embeddings")
async def get_document_embeddings(documento_id: UUID):
    """
    Retorna os embeddings de um documento
    """
    raise HTTPException(status_code=501, detail="Em desenvolvimento")


@router.post("/extract-image")
async def extract_image(
    file: UploadFile = File(...),
    tipo: str = Form("outro"),
):
    """Extract text from image using OCR (pytesseract)"""
    try:
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))

        # OCR extraction
        texto = pytesseract.image_to_string(image, lang='por')

        if not texto or not texto.strip():
            return {"texto_extraido": "", "dados_extraidos": {}, "erro": "Nenhum texto encontrado na imagem"}

        # Extract structured data if orcamento/apolice
        dados_extraidos = {}
        if tipo in ("orcamento", "apolice"):
            dados_extraidos = await _extract_structured_data(texto, tipo)

        return {
            "texto_extraido": texto,
            "dados_extraidos": dados_extraidos,
        }
    except Exception as e:
        logger.error(f"Erro ao extrair texto da imagem: {e}")
        return {"texto_extraido": "", "dados_extraidos": {}, "erro": str(e)}


@router.post("/extract-docx")
async def extract_docx(
    file: UploadFile = File(...),
    tipo: str = Form("outro"),
):
    """Extract text from DOCX file"""
    try:
        contents = await file.read()
        doc = Document(io.BytesIO(contents))

        # Extract all paragraphs
        texto = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    texto += "\n" + row_text

        if not texto or not texto.strip():
            return {"texto_extraido": "", "dados_extraidos": {}, "erro": "Nenhum texto encontrado no documento"}

        # Extract structured data if orcamento/apolice
        dados_extraidos = {}
        if tipo in ("orcamento", "apolice"):
            dados_extraidos = await _extract_structured_data(texto, tipo)

        return {
            "texto_extraido": texto,
            "dados_extraidos": dados_extraidos,
        }
    except Exception as e:
        logger.error(f"Erro ao extrair texto do DOCX: {e}")
        return {"texto_extraido": "", "dados_extraidos": {}, "erro": str(e)}
